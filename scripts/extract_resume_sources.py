from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = Path(__file__).resolve().parent / "resume_sources.json"


def extract_docx(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(path)
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {index}]\n{text}")
    return "\n\n".join(pages)


def main() -> None:
    candidates = sorted(
        path
        for path in ROOT.iterdir()
        if path.is_file()
        and not path.name.startswith("~$")
        and path.suffix.lower() in {".docx", ".pdf"}
    )
    payload = []
    for path in candidates:
        try:
            text = extract_docx(path) if path.suffix.lower() == ".docx" else extract_pdf(path)
            payload.append({"file": path.name, "text": text})
        except Exception as exc:
            payload.append({"file": path.name, "error": str(exc), "text": ""})
    OUTPUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Extracted {len(payload)} files to {OUTPUT}")


if __name__ == "__main__":
    main()
